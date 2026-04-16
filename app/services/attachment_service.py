from __future__ import annotations

import base64
from collections import deque
import csv
from io import BytesIO
from io import StringIO
import json
from pathlib import Path
import re

from bs4 import BeautifulSoup
try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional dependency fallback for local dev shells
    DocxDocument = None
try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional dependency fallback for local dev shells
    load_workbook = None
from PIL import Image, ImageChops, ImageEnhance, ImageFilter, ImageOps
from pypdf import PdfReader
from sqlalchemy.orm import Session
try:
    import xlrd
except Exception:  # pragma: no cover - optional dependency fallback for local dev shells
    xlrd = None

from app.core.config import get_settings
from app.models.uploaded_asset import UploadedAsset, UploadedAssetKind
from app.services.command_result import CommandResult, MessageAttachment
from app.services.llm_service import LLMService
from app.services.marketplace_search_service import MarketplaceSearchService
from app.services.upload_service import UploadService

settings = get_settings()


class AttachmentService:
    TEXT_EXTENSIONS = {
        ".txt",
        ".md",
        ".py",
        ".json",
        ".csv",
        ".tsv",
        ".log",
        ".yaml",
        ".yml",
        ".xml",
        ".html",
        ".htm",
        ".js",
        ".ts",
        ".tsx",
        ".jsx",
        ".ini",
        ".toml",
        ".sql",
        ".docx",
        ".docm",
        ".xlsx",
        ".xls",
        ".xlsm",
    }

    @staticmethod
    def build_message_attachments(assets: list[UploadedAsset]) -> list[MessageAttachment]:
        attachments: list[MessageAttachment] = []
        for asset in assets:
            attachments.append(
                MessageAttachment(
                    kind="photo" if asset.kind == UploadedAssetKind.image else "document",
                    path=asset.stored_path,
                    filename=asset.original_filename,
                    explicit_public_url=UploadService.build_public_url(asset),
                    asset_id=asset.asset_id,
                )
            )
        return attachments

    @staticmethod
    def build_metadata(assets: list[UploadedAsset]) -> list[dict]:
        return [
            {
                "asset_id": asset.asset_id,
                "kind": asset.kind.value,
                "filename": asset.original_filename,
                "mime_type": asset.mime_type,
                "size_bytes": asset.size_bytes,
                "public_url": UploadService.build_public_url(asset),
            }
            for asset in assets
        ]

    @staticmethod
    def try_handle(
        db: Session,
        user_id: str,
        session_id: str,
        text: str,
        assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult | None:
        if not assets:
            return None

        normalized_text = text.strip()
        image_assets = [asset for asset in assets if asset.kind == UploadedAssetKind.image]
        document_assets = [asset for asset in assets if asset.kind != UploadedAssetKind.image]

        if image_assets and AttachmentService._looks_like_edit_request(normalized_text):
            return AttachmentService._handle_image_edit(db, user_id, session_id, normalized_text, image_assets[0])

        if image_assets and AttachmentService._looks_like_marketplace_search_request(normalized_text):
            return AttachmentService._handle_image_marketplace_search(
                normalized_text=normalized_text,
                image_assets=image_assets,
                prompt_messages=prompt_messages,
                active_provider=active_provider,
                active_model=active_model,
            )

        if image_assets:
            return AttachmentService._handle_image_analysis(
                normalized_text=normalized_text,
                image_assets=image_assets,
                prompt_messages=prompt_messages,
                active_provider=active_provider,
                active_model=active_model,
            )

        if document_assets and AttachmentService._looks_like_document_edit_request(normalized_text):
            return AttachmentService._handle_document_edit(
                db=db,
                user_id=user_id,
                session_id=session_id,
                normalized_text=normalized_text,
                document_assets=document_assets,
                prompt_messages=prompt_messages,
                active_provider=active_provider,
                active_model=active_model,
            )

        if document_assets:
            return AttachmentService._handle_document_analysis(
                normalized_text=normalized_text,
                document_assets=document_assets,
                prompt_messages=prompt_messages,
                active_provider=active_provider,
                active_model=active_model,
            )

        return None

    @staticmethod
    def _handle_document_analysis(
        normalized_text: str,
        document_assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult:
        extracted_parts: list[str] = []
        for asset in document_assets[:4]:
            content = AttachmentService.extract_text(asset)
            extracted_parts.append(
                f"Attachment: {asset.original_filename}\n"
                f"MIME: {asset.mime_type}\n"
                f"Content excerpt:\n{content}"
            )

        request_text = normalized_text or "Summarize the uploaded file and highlight any important actions or risks."
        augmented_messages = list(prompt_messages)
        augmented_messages[-1] = {
            "role": "user",
            "content": f"{request_text}\n\nUploaded file context:\n\n" + "\n\n---\n\n".join(extracted_parts),
        }

        try:
            reply, provider = LLMService.generate_reply(
                augmented_messages,
                active_provider=active_provider,
                active_model=active_model,
            )
        except Exception as exc:
            _ = exc
            return CommandResult(
                reply=(
                    "I couldn't analyze the uploaded file right now because the AI model is temporarily busy. "
                    "Please try again in a minute."
                ),
                provider="attachment-error",
            )

        return CommandResult(
            reply=reply,
            provider=provider,
            attachments=AttachmentService.build_message_attachments(document_assets),
        )

    @staticmethod
    def _handle_document_edit(
        db: Session,
        user_id: str,
        session_id: str,
        normalized_text: str,
        document_assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult:
        source_asset = document_assets[0]
        source_text = AttachmentService.extract_text(source_asset)
        instruction = normalized_text or "Rewrite this file for clarity and keep the meaning intact."
        augmented_messages = list(prompt_messages)
        augmented_messages[-1] = {
            "role": "user",
            "content": (
                "Edit the uploaded file according to the instruction below.\n\n"
                f"Instruction: {instruction}\n\n"
                "Return only the fully edited file contents without commentary or markdown fencing.\n\n"
                f"Source filename: {source_asset.original_filename}\n\n"
                f"Source text:\n{source_text}"
            ),
        }

        try:
            reply, provider = LLMService.generate_reply(
                augmented_messages,
                active_provider=active_provider,
                active_model=active_model,
            )
        except Exception as exc:
            _ = exc
            return CommandResult(
                reply=(
                    "I couldn't edit the uploaded file right now because the AI model is temporarily busy. "
                    "Please try again in a minute."
                ),
                provider="attachment-error",
                attachments=AttachmentService.build_message_attachments(document_assets),
            )

        cleaned_reply = AttachmentService._strip_code_fences(reply).strip()
        suffix = Path(source_asset.original_filename).suffix.lower()
        output_suffix = suffix if suffix in AttachmentService.TEXT_EXTENSIONS else ".md"
        output_filename = f"{Path(source_asset.original_filename).stem}-edited{output_suffix}"
        mime_type = "text/markdown" if output_suffix == ".md" else "text/plain"
        generated = UploadService.create_generated_text_asset(
            db=db,
            platform_user_id=user_id,
            session_id=session_id,
            original_filename=output_filename,
            text=cleaned_reply,
            mime_type=mime_type,
            metadata_json={"source_asset_id": source_asset.asset_id, "instruction": instruction},
        )

        return CommandResult(
            reply="Edited file ready.",
            provider=provider,
            attachments=AttachmentService.build_message_attachments([generated]),
        )

    @staticmethod
    def _handle_image_analysis(
        normalized_text: str,
        image_assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult:
        if AttachmentService._should_offer_image_concierge(normalized_text):
            return AttachmentService._build_image_concierge(
                image_assets=image_assets,
                prompt_messages=prompt_messages,
                active_provider=active_provider,
                active_model=active_model,
            )

        if AttachmentService._looks_like_quick_identification_request(normalized_text):
            try:
                reply, provider = LLMService.generate_fast_vision_reply(
                    prompt_messages=prompt_messages,
                    prompt_text=(
                        "Identify what this image most likely shows in a fast, direct answer. "
                        "Mention any obvious text if it matters. Keep it short and practical."
                    ),
                    image_assets=image_assets,
                    active_provider=active_provider,
                    active_model=active_model,
                )
                return CommandResult(
                    reply=reply,
                    provider=provider,
                    attachments=AttachmentService.build_message_attachments(image_assets),
                )
            except Exception:
                pass

        request_text = normalized_text or "Describe this image clearly and extract any important text or visual details."
        if len(request_text.split()) <= 3 and any(token in request_text.lower() for token in ("verify", "identify", "what", "check")):
            request_text = (
                "Identify what this image shows, verify the important visual details, and answer the user's request clearly."
            )

        try:
            reply, provider = LLMService.generate_vision_reply(
                prompt_messages=prompt_messages,
                prompt_text=request_text,
                image_assets=image_assets,
                active_provider=active_provider,
                active_model=active_model,
            )
        except Exception as exc:
            _ = exc
            return CommandResult(
                reply=(
                    "I couldn't analyze the uploaded image right now because the AI model is temporarily busy. "
                    "Please try again in a minute."
                ),
                provider="attachment-error",
                attachments=AttachmentService.build_message_attachments(image_assets),
            )

        return CommandResult(
            reply=reply,
            provider=provider,
            attachments=AttachmentService.build_message_attachments(image_assets),
        )

    @staticmethod
    def _build_image_concierge(
        image_assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult:
        try:
            preview, provider = LLMService.generate_vision_reply(
                prompt_messages=prompt_messages,
                prompt_text=(
                    "Give a short one or two sentence description of the main subject in this image. "
                    "Mention visible text only if it is important. Do not use markdown bullets."
                ),
                image_assets=image_assets,
                active_provider=active_provider,
                active_model=active_model,
            )
        except Exception:
            preview = "I can see the uploaded image and it is ready for analysis."
            provider = "attachment-image"

        reply = (
            f"{preview.strip()}\n\n"
            "Tell me what you want me to do next. You can say:\n"
            "- `what is this`\n"
            "- `extract text`\n"
            "- `remove background`\n"
            "- `find Shopee links`\n"
            "- `find similar items`\n"
            "- `summarize the important details`"
        )
        return CommandResult(
            reply=reply,
            provider=provider,
            attachments=AttachmentService.build_message_attachments(image_assets),
        )

    @staticmethod
    def _handle_image_marketplace_search(
        normalized_text: str,
        image_assets: list[UploadedAsset],
        prompt_messages: list[dict[str, str]],
        active_provider: str,
        active_model: str,
    ) -> CommandResult:
        shopping_intent = (
            "Identify the main purchasable product in this image and return strict JSON only with these keys: "
            "`identified_item`, `search_query`, `confidence`, `key_attributes`, `notes`. "
            "`key_attributes` must be an array of short strings. "
            "Keep `search_query` concise and suitable for a shopping marketplace search."
        )

        provider = "attachment-marketplace"
        identified_item = ""
        search_query = ""
        confidence = ""
        key_attributes: list[str] = []
        notes = ""

        try:
            raw_reply, provider = LLMService.generate_vision_reply(
                prompt_messages=prompt_messages,
                prompt_text=shopping_intent,
                image_assets=image_assets,
                active_provider=active_provider,
                active_model=active_model,
            )
            parsed = AttachmentService._extract_json_dict(raw_reply)
            identified_item = str(parsed.get("identified_item", "")).strip()
            search_query = str(parsed.get("search_query", "")).strip()
            confidence = str(parsed.get("confidence", "")).strip()
            key_attributes = [
                str(item).strip()
                for item in (parsed.get("key_attributes") or [])
                if str(item).strip()
            ][:6]
            notes = str(parsed.get("notes", "")).strip()
        except Exception:
            try:
                raw_reply, provider = LLMService.generate_vision_reply(
                    prompt_messages=prompt_messages,
                    prompt_text=(
                        "Give the best shopping search keywords for this image in one short line only, "
                        "followed by a second short line that says what the item most likely is."
                    ),
                    image_assets=image_assets,
                    active_provider=active_provider,
                    active_model=active_model,
                )
                lines = [line.strip("- ").strip() for line in raw_reply.splitlines() if line.strip()]
                if lines:
                    search_query = lines[0]
                if len(lines) > 1:
                    identified_item = lines[1]
            except Exception:
                search_query = ""

        best_query = MarketplaceSearchService.normalize_query(
            search_query or identified_item or normalized_text,
            fallback="similar product",
        )
        links = MarketplaceSearchService.build_marketplace_links(best_query)

        lines = []
        if identified_item:
            lines.append(f"Best match: {identified_item}")
        else:
            lines.append("I built a best-effort shopping search from the uploaded image.")
        lines.append(f"Search query: `{best_query}`")
        if confidence:
            lines.append(f"Confidence: {confidence}")
        if key_attributes:
            lines.append("Key details: " + ", ".join(key_attributes))
        if notes:
            lines.append(notes)
        lines.extend(
            [
                "",
                "Marketplace links:",
            ]
        )
        for item in links:
            lines.append(f"- {item['label']}: {item['url']}")
        lines.extend(
            [
                "",
                "If you want, I can also `what is this`, `extract text`, or `remove background` for the same image.",
            ]
        )
        return CommandResult(
            reply="\n".join(lines),
            provider=provider,
            attachments=AttachmentService.build_message_attachments(image_assets),
        )

    @staticmethod
    def _handle_image_edit(
        db: Session,
        user_id: str,
        session_id: str,
        normalized_text: str,
        image_asset: UploadedAsset,
    ) -> CommandResult:
        image = Image.open(BytesIO(UploadService.load_bytes(image_asset))).convert("RGBA")
        operations: list[str] = []
        text_lower = normalized_text.lower()
        requested_background_removal = (
            "remove background" in text_lower
            or "background transparent" in text_lower
            or "transparent background" in text_lower
            or "cut out" in text_lower
            or "isolate subject" in text_lower
        )

        if "grayscale" in text_lower or "black and white" in text_lower:
            image = ImageOps.grayscale(image).convert("RGBA")
            operations.append("grayscale")
        if "invert" in text_lower:
            rgb = image.convert("RGB")
            image = ImageOps.invert(rgb).convert("RGBA")
            operations.append("invert")
        if "blur" in text_lower:
            image = image.filter(ImageFilter.GaussianBlur(radius=2.2))
            operations.append("blur")
        if "sharpen" in text_lower:
            image = image.filter(ImageFilter.SHARPEN)
            operations.append("sharpen")
        if "flip horizontal" in text_lower or "mirror" in text_lower:
            image = ImageOps.mirror(image)
            operations.append("flip horizontal")
        if "flip vertical" in text_lower:
            image = ImageOps.flip(image)
            operations.append("flip vertical")
        if "rotate left" in text_lower:
            image = image.rotate(90, expand=True)
            operations.append("rotate left")
        if "rotate right" in text_lower:
            image = image.rotate(-90, expand=True)
            operations.append("rotate right")
        if "rotate 180" in text_lower:
            image = image.rotate(180, expand=True)
            operations.append("rotate 180")
        if requested_background_removal:
            image, removed = AttachmentService._remove_background(image)
            if removed:
                operations.append("remove background")

        size_match = re.search(r"(\d{2,5})\s*[xX]\s*(\d{2,5})", normalized_text)
        if size_match:
            width = max(1, int(size_match.group(1)))
            height = max(1, int(size_match.group(2)))
            image = image.resize((width, height))
            operations.append(f"resize {width}x{height}")

        bright_match = re.search(r"bright(?:en|ness)?\s+(\d+)%", text_lower)
        if bright_match:
            factor = max(int(bright_match.group(1)), 1) / 100
            image = ImageEnhance.Brightness(image).enhance(factor)
            operations.append(f"brightness {bright_match.group(1)}%")

        contrast_match = re.search(r"contrast\s+(\d+)%", text_lower)
        if contrast_match:
            factor = max(int(contrast_match.group(1)), 1) / 100
            image = ImageEnhance.Contrast(image).enhance(factor)
            operations.append(f"contrast {contrast_match.group(1)}%")

        if not operations:
            if requested_background_removal:
                return CommandResult(
                    reply=(
                        "I couldn't confidently remove the background from that image. "
                        "Try a photo with a clearer subject separation or a simpler background."
                    ),
                    provider="attachment-edit",
                    attachments=AttachmentService.build_message_attachments([image_asset]),
                )
            return CommandResult(
                reply=(
                    "I can edit images with operations like `remove background`, `grayscale`, `blur`, "
                    "`sharpen`, `rotate left`, `flip horizontal`, or `resize 1024x1024`."
                ),
                provider="attachment-edit",
                attachments=AttachmentService.build_message_attachments([image_asset]),
            )

        output = BytesIO()
        image.save(output, format="PNG")
        output.seek(0)
        generated = UploadService.create_generated_image_asset(
            db=db,
            platform_user_id=user_id,
            session_id=session_id,
            original_filename=f"edited-{Path(image_asset.original_filename).stem}.png",
            image_bytes=output.getvalue(),
            metadata_json={"source_asset_id": image_asset.asset_id, "operations": operations},
        )

        return CommandResult(
            reply=f"Edited image ready. Applied: {', '.join(operations)}.",
            provider="attachment-edit",
            attachments=AttachmentService.build_message_attachments([generated]),
        )

    @staticmethod
    def extract_text(asset: UploadedAsset) -> str:
        raw = UploadService.load_bytes(asset)
        return AttachmentService.extract_text_from_bytes(
            original_filename=asset.original_filename,
            mime_type=asset.mime_type,
            raw=raw,
        )

    @staticmethod
    def extract_text_from_bytes(original_filename: str, mime_type: str | None, raw: bytes) -> str:
        suffix = Path(original_filename).suffix.lower()
        normalized_mime = (mime_type or "").lower()

        if suffix == ".pdf" or normalized_mime == "application/pdf":
            text = AttachmentService._extract_pdf_text(raw)
        elif suffix in {".docx", ".docm"} or normalized_mime in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }:
            text = AttachmentService._extract_docx_text(raw)
        elif suffix == ".doc":
            text = (
                "Legacy `.doc` files are not fully supported by the built-in parser yet. "
                "Please re-save the document as `.docx` for best extraction quality."
            )
        elif suffix in {".xlsx", ".xlsm"} or normalized_mime in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel.sheet.macroenabled.12",
        }:
            text = AttachmentService._extract_xlsx_text(raw)
        elif suffix == ".xls" or normalized_mime == "application/vnd.ms-excel":
            text = AttachmentService._extract_xls_text(raw)
        elif suffix in {".csv", ".tsv"}:
            text = AttachmentService._extract_delimited_text(raw, delimiter="\t" if suffix == ".tsv" else ",")
        elif suffix in {".html", ".htm"} or "html" in normalized_mime:
            text = AttachmentService._extract_html_text(raw)
        else:
            text = AttachmentService._decode_text_bytes(raw)

        text = text.strip()
        if not text:
            text = "(No readable text content found.)"
        return text[: settings.upload_extract_text_chars]

    @staticmethod
    def encode_image_asset(asset: UploadedAsset) -> tuple[str, str]:
        raw = UploadService.load_bytes(asset)
        return base64.b64encode(raw).decode("ascii"), asset.mime_type

    @staticmethod
    def _should_offer_image_concierge(text: str) -> bool:
        lowered = (text or "").strip().lower()
        return lowered in {
            "",
            "image",
            "photo",
            "picture",
            "look at this",
            "check this out",
            "see this",
        }

    @staticmethod
    def _looks_like_quick_identification_request(text: str) -> bool:
        lowered = (text or "").strip().lower()
        if not lowered:
            return False
        quick_tokens = (
            "what is this",
            "what's this",
            "tell me what this is",
            "can you tell me what this is",
            "identify this",
            "identify it",
            "recognize this",
            "recognise this",
            "what is in this image",
            "what is in this photo",
            "what is in this picture",
        )
        return any(token in lowered for token in quick_tokens)

    @staticmethod
    def should_use_recent_assets(text: str) -> bool:
        lowered = text.lower().strip()
        if not lowered:
            return False
        direct_tokens = (
            "this image",
            "this photo",
            "this picture",
            "this screenshot",
            "this file",
            "this document",
            "this pdf",
            "attached image",
            "attached file",
            "uploaded image",
            "uploaded file",
            "last image",
            "last file",
            "previous image",
            "previous file",
            "the image",
            "the file",
            "this attachment",
            "the attachment",
            "the upload",
            "that image",
            "that photo",
            "that picture",
            "that file",
        )
        attachment_intents = (
            "what is this",
            "what's this",
            "identify this",
            "identify it",
            "recognize this",
            "recognise this",
            "verify this",
            "check this image",
            "check this file",
            "check it",
            "check this photo",
            "check this picture",
            "describe this",
            "describe it",
            "summarize this",
            "summarise this",
            "summarize it",
            "summarise it",
            "analyze this",
            "analyse this",
            "analyze it",
            "analyse it",
            "read this image",
            "read the image",
            "read this file",
            "read that file",
            "extract text",
            "ocr this",
            "edit it",
            "edit this",
            "remove bg",
            "remove the background",
            "remove background",
            "find shopee links",
            "find lazada links",
            "where can i buy this",
            "where to buy this",
            "buy this",
            "find similar items",
            "price check",
            "compare price",
            "shopping links",
            "seller link",
        )
        return (
            any(token in lowered for token in direct_tokens)
            or any(token in lowered for token in attachment_intents)
            or AttachmentService._looks_like_marketplace_search_request(lowered)
            or AttachmentService._looks_like_edit_request(lowered)
            or AttachmentService._looks_like_document_edit_request(lowered)
        )

    @staticmethod
    def _looks_like_marketplace_search_request(text: str) -> bool:
        return MarketplaceSearchService.looks_like_marketplace_request(text)

    @staticmethod
    def _looks_like_edit_request(text: str) -> bool:
        lowered = text.lower()
        return any(
            token in lowered
            for token in (
                "edit",
                "grayscale",
                "black and white",
                "blur",
                "sharpen",
                "rotate",
                "flip",
                "resize",
                "contrast",
                "brightness",
                "invert",
                "crop",
                "background",
                "transparent",
                "cut out",
                "isolate subject",
                "remove bg",
            )
        )

    @staticmethod
    def _looks_like_document_edit_request(text: str) -> bool:
        lowered = text.lower()
        return any(
            token in lowered
            for token in (
                "rewrite",
                "edit this file",
                "edit this document",
                "edit this text",
                "make this shorter",
                "make this clearer",
                "fix grammar",
                "improve writing",
                "reformat",
                "translate",
                "convert this",
                "clean up this file",
                "refactor this text",
            )
        )

    @staticmethod
    def _strip_code_fences(text: str) -> str:
        cleaned = text.strip()
        if cleaned.startswith("```") and cleaned.endswith("```"):
            lines = cleaned.splitlines()
            if len(lines) >= 3:
                return "\n".join(lines[1:-1])
        return cleaned

    @staticmethod
    def _extract_json_dict(text: str) -> dict:
        cleaned = AttachmentService._strip_code_fences(text).strip()
        if cleaned.startswith("json"):
            cleaned = cleaned[4:].strip()
        try:
            parsed = json.loads(cleaned)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            pass

        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            try:
                parsed = json.loads(cleaned[start : end + 1])
                return parsed if isinstance(parsed, dict) else {}
            except json.JSONDecodeError:
                return {}
        return {}

    @staticmethod
    def _decode_text_bytes(raw: bytes) -> str:
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("utf-8", errors="replace")

    @staticmethod
    def _extract_pdf_text(raw: bytes) -> str:
        reader = PdfReader(BytesIO(raw))
        parts = []
        for page in reader.pages[:20]:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)

    @staticmethod
    def _extract_docx_text(raw: bytes) -> str:
        if DocxDocument is None:
            raise RuntimeError("DOCX extraction dependency is not installed.")
        document = DocxDocument(BytesIO(raw))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables[:8]:
            for row in table.rows[:30]:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    parts.append(" | ".join(value for value in values if value))
        return "\n".join(parts)

    @staticmethod
    def _extract_xlsx_text(raw: bytes) -> str:
        if load_workbook is None:
            raise RuntimeError("XLSX extraction dependency is not installed.")
        workbook = load_workbook(filename=BytesIO(raw), read_only=True, data_only=True)
        parts: list[str] = []
        try:
            for sheet in workbook.worksheets[:6]:
                rows: list[str] = []
                for row in sheet.iter_rows(min_row=1, max_row=40, values_only=True):
                    values = ["" if value is None else str(value) for value in row[:20]]
                    if any(value.strip() for value in values):
                        rows.append("\t".join(values).rstrip())
                if rows:
                    parts.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
        finally:
            workbook.close()
        return "\n\n".join(parts)

    @staticmethod
    def _extract_xls_text(raw: bytes) -> str:
        if xlrd is None:
            raise RuntimeError("XLS extraction dependency is not installed.")
        workbook = xlrd.open_workbook(file_contents=raw)
        parts: list[str] = []
        for sheet in workbook.sheets()[:6]:
            rows: list[str] = []
            for row_idx in range(min(sheet.nrows, 40)):
                values = [str(sheet.cell_value(row_idx, col_idx)).strip() for col_idx in range(min(sheet.ncols, 20))]
                if any(values):
                    rows.append("\t".join(values).rstrip())
            if rows:
                parts.append(f"Sheet: {sheet.name}\n" + "\n".join(rows))
        return "\n\n".join(parts)

    @staticmethod
    def _extract_delimited_text(raw: bytes, delimiter: str) -> str:
        text = AttachmentService._decode_text_bytes(raw)
        reader = csv.reader(StringIO(text), delimiter=delimiter)
        rows: list[str] = []
        for index, row in enumerate(reader):
            if index >= 60:
                break
            rows.append("\t".join(str(cell).strip() for cell in row))
        return "\n".join(rows)

    @staticmethod
    def _extract_html_text(raw: bytes) -> str:
        decoded = AttachmentService._decode_text_bytes(raw)
        soup = BeautifulSoup(decoded, "html.parser")
        for tag_name in ("script", "style", "noscript", "svg", "canvas"):
            for tag in soup.find_all(tag_name):
                tag.decompose()

        parts: list[str] = []
        if soup.title and soup.title.string:
            parts.append(f"Title: {soup.title.string.strip()}")
        description = soup.find("meta", attrs={"name": re.compile("^description$", re.IGNORECASE)})
        if description and description.get("content"):
            parts.append(f"Description: {description.get('content', '').strip()}")

        body_text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
        if body_text:
            parts.append(body_text)
        return "\n\n".join(parts)

    @staticmethod
    def _remove_background(image: Image.Image) -> tuple[Image.Image, bool]:
        rgba = image.convert("RGBA")
        width, height = rgba.size
        if width < 2 or height < 2:
            return rgba, False

        pixels = rgba.load()
        border_points: list[tuple[int, int]] = []
        for x in range(width):
            border_points.append((x, 0))
            border_points.append((x, height - 1))
        for y in range(1, height - 1):
            border_points.append((0, y))
            border_points.append((width - 1, y))

        border_colors = [pixels[x, y][:3] for x, y in border_points if pixels[x, y][3] > 0]
        if not border_colors:
            return rgba, False

        background_color = tuple(int(round(sum(color[index] for color in border_colors) / len(border_colors))) for index in range(3))

        def color_distance(left: tuple[int, int, int], right: tuple[int, int, int]) -> int:
            return abs(left[0] - right[0]) + abs(left[1] - right[1]) + abs(left[2] - right[2])

        distances = sorted(color_distance(color, background_color) for color in border_colors)
        base_threshold = distances[int(len(distances) * 0.85)] if distances else 48
        threshold = max(30, min(120, base_threshold + 18))
        soften_threshold = min(180, threshold + 42)

        visited = bytearray(width * height)
        queue: deque[tuple[int, int]] = deque()

        def index_for(x: int, y: int) -> int:
            return y * width + x

        def qualifies(x: int, y: int, max_distance: int) -> bool:
            red, green, blue, alpha = pixels[x, y]
            if alpha == 0:
                return True
            return color_distance((red, green, blue), background_color) <= max_distance

        for x, y in border_points:
            idx = index_for(x, y)
            if visited[idx]:
                continue
            if qualifies(x, y, threshold):
                visited[idx] = 1
                queue.append((x, y))

        removed_pixels = 0
        while queue:
            x, y = queue.popleft()
            removed_pixels += 1
            for next_x, next_y in ((x - 1, y), (x + 1, y), (x, y - 1), (x, y + 1)):
                if next_x < 0 or next_y < 0 or next_x >= width or next_y >= height:
                    continue
                idx = index_for(next_x, next_y)
                if visited[idx]:
                    continue
                if qualifies(next_x, next_y, threshold):
                    visited[idx] = 1
                    queue.append((next_x, next_y))

        if removed_pixels == 0 or removed_pixels >= int(width * height * 0.98):
            return rgba, False

        mask = Image.new("L", (width, height), 255)
        mask_pixels = mask.load()
        for y in range(height):
            for x in range(width):
                idx = index_for(x, y)
                if visited[idx]:
                    mask_pixels[x, y] = 0
                    continue
                if qualifies(x, y, soften_threshold):
                    red, green, blue, _ = pixels[x, y]
                    distance = color_distance((red, green, blue), background_color)
                    fade = int(
                        max(
                            0,
                            min(
                                255,
                                255 * (distance - threshold) / max(1, soften_threshold - threshold),
                            ),
                        )
                    )
                    mask_pixels[x, y] = min(mask_pixels[x, y], fade)

        softened_mask = mask.filter(ImageFilter.GaussianBlur(radius=max(1.2, min(width, height) / 320)))
        result = rgba.copy()
        result.putalpha(ImageChops.multiply(rgba.getchannel("A"), softened_mask))
        return result, True
