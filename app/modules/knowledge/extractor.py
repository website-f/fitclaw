"""Plain-text extraction for supported KB document kinds."""
from __future__ import annotations

import csv
import io
from pathlib import Path

SUPPORTED_KINDS = {"pdf", "docx", "txt", "md", "csv"}


def detect_kind(filename: str, content_type: str | None = None) -> str:
    name = filename.lower()
    if name.endswith(".pdf") or (content_type or "").endswith("pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".csv"):
        return "csv"
    if name.endswith(".md") or name.endswith(".markdown"):
        return "md"
    return "txt"


def extract_text(*, data: bytes, kind: str, filename: str) -> str:
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "csv":
        return _extract_csv(data)
    return _extract_plain(data)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(f"[Page {page_no}]\n{text.strip()}")
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return ""
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows: list[str] = []
    for row in reader:
        cleaned = [cell.strip() for cell in row if cell and cell.strip()]
        if cleaned:
            rows.append(" | ".join(cleaned))
    return "\n".join(rows)


def _extract_plain(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def derive_title(filename: str, fallback: str = "Untitled document") -> str:
    stem = Path(filename).stem.strip()
    if not stem:
        return fallback
    # Replace underscores/hyphens with spaces for nicer titles.
    return stem.replace("_", " ").replace("-", " ").strip().title() or fallback
